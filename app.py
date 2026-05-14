from utils.file_loader import extract_text
from utils.summary_tools import generate_summary
from utils.resume_tools import resume_analyzer
import gradio as gr
from pypdf import PdfReader
import docx2txt
from langchain_community.chat_models import ChatOllama

MODEL_NAME = "llama3.2:3b"


def clean_text(text):
    text = text.replace("\n", " ")
    text = " ".join(text.split())
    return text

def calculate_ats_score(resume_text, job_text):

    resume_text = resume_text.lower()
    job_text = job_text.lower()

    keywords = list(set(job_text.split()))

    matched = []

    for word in keywords:
        if len(word) > 4 and word in resume_text:
            matched.append(word)

    if len(keywords) == 0:
        score = 0
    else:
        score = int((len(matched) / len(keywords)) * 100)

    score = min(score, 95)

    return score, matched[:25]

def ask_ai(prompt):
    llm = ChatOllama(
        model=MODEL_NAME,
        temperature=0.2
    )
    response = llm.invoke(prompt)
    return response.content


def job_matching(text, job_description):
    ats_score, matched_keywords = calculate_ats_score(text, job_description)
    prompt = f"""
You are PrimeAI, a professional AI job matching and career advisor.

Compare the uploaded resume with the job description below.

# PrimeAI ATS Resume Report
You MUST return the report using these exact headings only.

## ATS Score
Give a realistic ATS score out of 100.

## Best-Fit Roles
List 4 realistic roles:
- AI Trainer
- Data Analyst
- AI Operations Support
- Internal AI Tools Assistant

## Matched Keywords
List matched resume keywords.

## Missing Keywords
List missing or weak AI/data keywords.

## Key Strengths
List 3 resume strengths that match the job.

## Areas for Improvement
List 3 practical ATS improvements.

## Rewrite Suggestions
Rewrite weak resume bullet points professionally.

## Final Recommendation
Give a short hiring-positioning recommendation.
- AI Trainer
- Data Analyst
- AI Operations
- QA / AI Testing
- Internal AI Tools
- Technical Support AI
- Business Intelligence
- Analytics Support
- AI Workflow Support
- Research Support
# Final Recommendation
Explain whether the candidate should apply and what to improve.
Formatting rules:
- Use professional markdown headings.
- Use bullet points for strengths and weaknesses.
- Keep answers concise and ATS-focused.
- Sound like a professional AI career assistant.
- Avoid generic AI buzzwords.
- Prioritize realistic entry-to-intermediate AI/data roles.


Resume:
{text}

Job Description:
{job_description}
"""
    ai_response = ask_ai(prompt)

    formatted_report = f"""
# PrimeAI ATS Resume Report

## ATS Score
**{ats_score}/100**

## Best-Fit Roles
- AI Trainer
- Data Analyst
- AI Operations Support
- Internal AI Tools Assistant

## Matched Keywords
{matched_keywords}

## AI Resume Analysis
{ai_response}

## Final Recommendation
Use this resume for AI Trainer, Data Analyst, AI Operations, and Internal AI Tools roles after adding stronger measurable achievements, project links, GitHub links, and Hugging Face live demo links.
"""

    return formatted_report


def document_summary(text, question):
    cleaned_text = clean_text(text)

    prompt = f"""
You are PrimeAI Document QA Assistant.

Answer the user's question ONLY using information from the uploaded document.

Important rules:
- Stay grounded to the uploaded document.
- Do NOT invent information.
- Do NOT generate career coaching unless specifically asked.
- Keep answers concise, professional, and easy to read.
- Avoid repeating the same phrases.
- Avoid buzzwords and exaggerated AI wording.
- Use natural recruiter-style or business-style language.
- Double-check technology names and spelling carefully.

Formatting rules:
- Use short headings when useful.
- Use bullet points for lists.
- Keep paragraphs short.
- Avoid overly long explanations.

If the user asks for:
- a summary → provide a concise summary of the main points
- skills → list only the important skills found in the document
- jobs/roles → recommend only realistic roles supported by the document
- certifications → list certifications mentioned
- technologies/tools → list tools and technologies clearly

If the answer is not found in the uploaded document, say:
"I could not find that information in the uploaded document."

Uploaded Document:

{cleaned_text}

User Question:

{question}
"""
    return ask_ai(prompt)

def respond(message, history, file, mode):
    print("SELECTED MODE:", mode)
    if not message:
        return "Please type a question."

    text = extract_text(file)

    if text == "No file uploaded.":
        return "Please upload a PDF or Word document first."

    if mode == "Resume Analyzer + ATS":
        return resume_analyzer(text, message, ask_ai)

    if mode == "Job Matching":
        return job_matching(text, message)

    if mode == "Document QA / Summary":
        return document_summary(text, message)

    return "Please select a mode."

def create_ats_reports(answer):
    import os
    import tempfile
    from datetime import datetime

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = os.path.join(tempfile.gettempdir(), f"PrimeAI_ATS_Report_{timestamp}.pdf")
    word_path = os.path.join(tempfile.gettempdir(), f"PrimeAI_ATS_Report_{timestamp}.docx")

    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(pdf_path)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("PrimeAI ATS Resume Report", styles["Title"]))
        story.append(Spacer(1, 12))

        for line in answer.split("\n"):
            if line.strip():
                clean_line = line.replace("#", "").replace("*", "")
                story.append(Paragraph(clean_line, styles["BodyText"]))
                story.append(Spacer(1, 6))

        doc.build(story)
    except Exception:
        pdf_path = None

    try:
        from docx import Document

        document = Document()
        document.add_heading("PrimeAI ATS Resume Report", 0)

        for line in answer.split("\n"):
            clean_line = line.replace("#", "").replace("*", "").strip()
            if clean_line:
                document.add_paragraph(clean_line)

        document.save(word_path)
    except Exception:
        word_path = None

    return pdf_path, word_path

def chat_response(message, history, file, mode):

    if history is None:
        history = []

    if not message:
        return "", history, None, None

    try:
        answer = respond(message, history, file, mode)

        pdf_path, word_path = create_ats_reports(answer)

        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": answer})

        return "", history, pdf_path, word_path
    except Exception as e:
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": f"Error: {str(e)}"})

        return "", history, pdf_path, word_path


custom_css = """
body {
    background: #0b1220;
}

.gradio-container {
    max-width: 100% !important;
    width: 1400px !important;
    margin: 0 auto !important;
    padding: 12px !important;
}

#workspace {
    display: flex !important;
    flex-direction: row !important;
    justify-content: center !important;
    gap: 20px !important;
    align-items: flex-start !important;
}

#sidebar {
    min-width: 320px !important;
    max-width: 340px !important;
    height: fit-content !important;
}

#header {
    text-align: center !important;
    width: 100% !important;
}

#header h1 {
    margin-bottom: 8px !important;
}

#header p {
    font-size: 14px !important;
}

#header {
    text-align: center;
    padding: 28px;
    border-radius: 18px;
    background: linear-gradient(135deg, #111827, #1e293b);
    border: 1px solid #334155;
    margin-bottom: 18px;
}

#header h1 {
    color: #fde68a;
    font-size: 34px;
    margin-bottom: 8px;
}

#header p {
    color: #e5e7eb;
    font-size: 16px;
}

#sidebar {
    background: #111827;
    padding: 18px;
    border-radius: 18px;
    border: 1px solid #334155;
}

#main-area {
    flex: 1 !important;
    min-width: 0 !important;

    background: #111827;
    padding: 22px;
    border-radius: 20px;
    border: 1px solid #334155;

    box-shadow: 0 0 18px rgba(0,0,0,0.25);

    max-height: 92vh;
    overflow-y: auto;
}

.gr-button {
    font-weight: bold !important;
    border-radius: 12px !important;
}
.gradio-container {
    max-width: 1450px !important;
    margin: auto !important;
}



#header {
    padding: 12px 18px !important;
    margin-bottom: 20px !important;
}

textarea {
    min-height: 70px !important;

    background: #111827 !important;
    color: #f8fafc !important;

    border-radius: 14px !important;
    border: 1px solid #334155 !important;

    padding: 14px !important;

    font-size: 15px !important;
    line-height: 1.6 !important;

    box-shadow: none !important;
}
textarea:focus {
    border: 1px solid #60a5fa !important;
    box-shadow: 0 0 10px rgba(96,165,250,0.25) !important;
}

.gr-chatbot {
    height: 560px !important;

    background: #0f172a !important;

    border-radius: 18px !important;
    border: 1px solid #334155 !important;

    padding: 10px !important;

    font-size: 15px !important;
    line-height: 1.7 !important;

    overflow-y: auto !important;
}

.message {
    padding: 14px !important;
    border-radius: 14px !important;

    margin-bottom: 12px !important;

    line-height: 1.7 !important;
    font-size: 15px !important;
}

.gr-button {
    border-radius: 10px !important;
}

.gr-box {
    border-radius: 14px !important;
}

footer {
    display: none !important;
}
"""

with gr.Blocks(title="PrimeAI", css=custom_css) as demo:

    gr.HTML(
        """
        <div id="header">
            <h1>🤖 PrimeAI</h1>
            <p>
            Private AI Workspace — Document QA, Resume Analyzer,
            Job Matching, and ATS
            </p>
        </div>
        """
    )

    with gr.Row(elem_id="workspace"):
        with gr.Column(scale=1, elem_id="sidebar"):
            file_input = gr.File(
                label="Upload PDF or Word Document",
                file_types=[".pdf", ".docx"],
                height=140
            )

            mode_input = gr.Dropdown(
                choices=[
                    "Document QA / Summary",
                    "Resume Analyzer + ATS",
                    "Job Matching"
                ],
                value="Document QA / Summary",
                label="PrimeAI Mode"
            )

            clear_btn = gr.Button("Clear Chat")

        with gr.Column(scale=3, elem_id="main-area"):
            chatbot = gr.Chatbot(
                label="PrimeAI Chat",
                height=420,
                render_markdown=True
            )

            message_input = gr.Textbox(
                label="Ask PrimeAI",
                placeholder="Ask a question, analyze a resume, or paste a job description...",
                lines=2
            )

            ask_btn = gr.Button("Ask PrimeAI")
            pdf_output = gr.File(label="Download PDF ATS Report", interactive=False)
            word_output = gr.File(label="Download Word ATS Report", interactive=False)

    ask_btn.click(
        fn=chat_response,
        inputs=[message_input, chatbot, file_input, mode_input],
        outputs=[message_input, chatbot, pdf_output, word_output]
    )

    message_input.submit(
        fn=chat_response,
        inputs=[message_input, chatbot, file_input, mode_input],
        outputs=[message_input, chatbot, pdf_output, word_output]
    )

    clear_btn.click(lambda: [], outputs=chatbot)


demo.launch(server_port=7865, css=custom_css)